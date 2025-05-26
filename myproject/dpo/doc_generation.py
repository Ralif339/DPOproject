import docx
from .models import *
import io
from django.http import HttpResponse
from docx.shared import Pt
from transliterate import translit
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import locale
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

locale.setlocale(locale.LC_TIME, 'Russian')

def set_font(paragraph, pt: float, ):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(pt)
        
def replace_text(paragraph, old_text: str, new_text: str, pt: float):
    paragraph.text = paragraph.text.replace(old_text, new_text)
    set_font(paragraph, pt)
    
def sanitize_filename(filename):
    # Транслитерируем кириллицу в латиницу
    filename = translit(filename, 'ru', reversed=True)
    # Заменяем проблемные символы на _
    filename = ''.join(c if c.isalnum() or c in (' ', '.') else '_' for c in filename)
    return filename

def set_cell_borders(cell):
    """
    Устанавливает границы для ячейки.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Создаем элемент границ
    tcBorders = OxmlElement('w:tcBorders')

    # Добавляем границы (верхняя, нижняя, левая, правая)
    for border_type in ['top', 'bottom', 'left', 'right']:
        border = OxmlElement(f'w:{border_type}')
        border.set(qn('w:sz'), '5')  # Толщина границы (в восьмых частях точки)
        border.set(qn('w:val'), 'single')  # Тип линии (single, double и т.д.)
        border.set(qn('w:color'), '000000')  # Цвет границы (черный)
        tcBorders.append(border)

    # Добавляем границы в свойства ячейки
    tcPr.append(tcBorders)

def get_enroll_docx(group_id, number, date):
    group = Group.objects.get(id=group_id)
    expelled_students = StudentExpulsion.objects.filter(group=group).exclude(reason="Завершение курса").values_list("student_id", flat=True)
    student_groups = StudentGroup.objects.filter(group=group).exclude(student_id__in=expelled_students).select_related("student").order_by("student__surname")
    document = docx.Document('dpo/docs_patterns/enroll.docx')
    table1 = document.tables[0]
    table1.cell(1, 0).text = number
    table1.cell(1, 1).text = date
    for cell in table1.rows[1].cells:
        for run in cell.paragraphs[0].runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.paragraphs[9]
    match group.course.course_type:
        case "Профессиональная переподготовка":
            replace_text(p, "<CourseType>", "программе профессиональной подготовки", 11)
        case "Курсы повышения квалификации КПК":
            replace_text(p, "<CourseType>", "программе повышения квалификации КПК", 11)            
        case "Общеобразовательные программы для детей и взрослых":
            replace_text(p, "<CourseType>", "общеобразовательной программе для детей и взрослых", 11)            
        case "Профессиональное обучение":
            replace_text(p, "<CourseType>", "программе профессионального обучения", 11)
    replace_text(p, "<CourseName>", group.course.course_name, 11)
    p = document.paragraphs[11]
    replace_text(p, "<GroupName>", group.name, 12)
    for run in p.runs:
        run.bold = True
    p = document.paragraphs[12]
    replace_text(p, "<Cost>", f"{group.course.price_b} руб., {group.course.price_vb}", 11)
    
    table2 = document.tables[1]
    p = document.add_paragraph("(бюджет, внебюджет)")
    table2._element.addprevious(p._p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(11)
    
    p = document.paragraphs[15]
    replace_text(p, "<TeacherName>", f"{group.teacher.surname} {group.teacher.name[0]}. {group.teacher.patronymic[0]}.", 11)

    i = 1
    for student_group in student_groups:
        row = table2.add_row()
        row.cells[0].text = f"{i}."
        row.cells[1].text = f"{student_group.student.surname} {student_group.student.name} {student_group.student.patronymic}"
        if student_group.ed_kind == "Внебюджет":
            row.cells[2].text = "ВБС"
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
        i += 1
    
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    safe_group_name = sanitize_filename(group.name) 
    
    response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="enroll_order_{safe_group_name}.docx"'
    
    return response


def get_commission_docx(group_id, number, date):
    group = Group.objects.get(id=group_id)
    expelled_students = StudentExpulsion.objects.filter(group=group).exclude(reason="Завершение курса").values_list("student_id", flat=True)
    student_groups = StudentGroup.objects.filter(group=group).exclude(student_id__in=expelled_students).select_related("student").order_by("student__surname")
    document = docx.Document('dpo/docs_patterns/commission.docx')
    
    table1 = document.tables[0]
    table1.cell(1, 0).text = number
    table1.cell(1, 1).text = date
    for cell in table1.rows[1].cells:
        for run in cell.paragraphs[0].runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    p = document.paragraphs[9]
    match group.course.course_type:
        case "Профессиональная переподготовка":
            replace_text(p, "<CourseType>", "программе профессиональной подготовки", 11)
        case "Курсы повышения квалификации КПК":
            replace_text(p, "<CourseType>", "программе повышения квалификации КПК", 11)            
        case "Общеобразовательные программы для детей и взрослых":
            replace_text(p, "<CourseType>", "общеобразовательной программе для детей и взрослых", 11)            
        case "Профессиональное обучение":
            replace_text(p, "<CourseType>", "программе профессионального обучения", 11)
    replace_text(p, "<CourseName>", group.course.course_name, 11)
    
    for group_commission in GroupCommission.objects.filter(group=group):
        match group_commission.role:
            case "Председатель комиссии":
                p = document.paragraphs[10]
                chairmanFIO = f"{group_commission.commission_member.surname} {group_commission.commission_member.name[0]}. {group_commission.commission_member.patronymic[0]}."
                replace_text(p, "<ChairmanFIO>", chairmanFIO, 11)
                replace_text(p, "<ChairmanPosition>", group_commission.commission_member.position, 11)
            case "Заместитель председателя комиссии":
                p = document.paragraphs[11]
                deputyFIO = f"{group_commission.commission_member.surname} {group_commission.commission_member.name[0]}. {group_commission.commission_member.patronymic[0]}."
                replace_text(p, "<DeputyFIO>", deputyFIO, 11)
                replace_text(p, "<DeputyPosition>", group_commission.commission_member.position, 11)
            case "Член комиссии":
                p = document.paragraphs[12]
                memberFIO = f"{group_commission.commission_member.surname} {group_commission.commission_member.name[0]}. {group_commission.commission_member.patronymic[0]}."
                replace_text(p, "<MemberFIO>", memberFIO, 11)
                replace_text(p, "<MemberPosition>", group_commission.commission_member.position, 11)
            case "Секретарь":
                p = document.paragraphs[13]
                secretaryFIO = f"{group_commission.commission_member.surname} {group_commission.commission_member.name[0]}. {group_commission.commission_member.patronymic[0]}."
                replace_text(p, "<SecretaryFIO>", secretaryFIO, 11)
                replace_text(p, "<SecretaryPosition>", group_commission.commission_member.position, 11)
                
    p = document.paragraphs[15]
    replace_text(p, "<GroupName>", group.name, 12)
    for run in p.runs:
        run.bold = True
        
    p = document.paragraphs[16]
    replace_text(p, "<Price>", f"{group.course.price_b} руб., {group.course.price_vb}", 11)

    table2 = document.tables[1]
    i = 1
    for student_group in student_groups:
        row = table2.add_row()
        row.cells[0].text = f"{i}."
        row.cells[1].text = f"{student_group.student.surname} {student_group.student.name} {student_group.student.patronymic}"
        if student_group.ed_kind == "Внебюджет":
            row.cells[2].text = "ВБС"
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
        i += 1
    
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    safe_group_name = sanitize_filename(group.name) 
    
    response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="commission_order_{safe_group_name}.docx"'
    
    return response


def get_lesson_log_docx(group_id):
    group = Group.objects.get(id=group_id)
    expelled_students = StudentExpulsion.objects.filter(group=group).exclude(reason="Завершение курса").values_list("student_id", flat=True)
    student_groups = StudentGroup.objects.filter(group=group).exclude(student_id__in=expelled_students).select_related("student").order_by("student__surname")
    document = docx.Document('dpo/docs_patterns/lesson_log.docx')
    
    p = document.paragraphs[0]
    replace_text(p, "<CourseName>", f"{group.course.course_name}", 11)
    for run in p.runs:
        run.font.name = "GOST type A"
        run.italic = True
        
    p = document.paragraphs[2]
    replace_text(p, "<TeacherFIO>", f"{group.teacher.surname} {group.teacher.name[0]}. {group.teacher.patronymic[0]}.", 12)
    for run in p.runs:
        run.font.name = "GOST type A"
        run.italic = True
        
    table1 = document.tables[0]
    table1.rows[2].cells[0].text = group.name
    table1.rows[2].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    students = []
    if student_groups:
        students = [f"{student_group.student.surname} {student_group.student.name[0]}. {student_group.student.patronymic[0]}."
                    for student_group in student_groups]
    
        students.sort(reverse=True)

        for row in table1.rows[3:]:
            while students:
                row.cells[1].text = students.pop()
        
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    safe_group_name = sanitize_filename(group.name) 
    
    response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="lesson_log_{safe_group_name}.docx"'
    
    return response

def get_protocol_docx(group_id, doc_number, date):
    group = Group.objects.get(id=group_id)
    expelled_students = StudentExpulsion.objects.filter(group=group).exclude(reason="Завершение курса").values_list("student_id", flat=True)
    student_groups = StudentGroup.objects.filter(group=group).exclude(student_id__in=expelled_students).select_related("student").order_by("student__surname")
    document = docx.Document('dpo/docs_patterns/protocol.docx')
    # Преобразуем строку в объект datetime
    date_obj = datetime.strptime(date, '%Y-%m-%d')

    # Форматируем дату в нужный формат
    formatted_date = date_obj.strftime('«%d» %B %Y г.')
    
    p = document.paragraphs[17]
    replace_text(p, "<Number>", f"{doc_number}", 11)
    replace_text(p, "<Date>", f"{formatted_date}", 11)
    
    for group_commission in GroupCommission.objects.filter(group=group):
        match group_commission.role:
            case "Председатель комиссии":
                p = document.paragraphs[22]
                chairmanFIO = f"{group_commission.commission_member.surname} {group_commission.commission_member.name[0]}. {group_commission.commission_member.patronymic[0]}."
                replace_text(p, "<ChairmanFIO>", chairmanFIO, 11)
                replace_text(p, "<ChairmanPosition>", group_commission.commission_member.position, 11)
                p = document.paragraphs[32]
                replace_text(p, "<ChairmanFIO>", chairmanFIO, 11)
            case "Заместитель председателя комиссии":
                p = document.paragraphs[23]
                deputyFIO = f"{group_commission.commission_member.surname} {group_commission.commission_member.name[0]}. {group_commission.commission_member.patronymic[0]}."
                replace_text(p, "<DeputyFIO>", deputyFIO, 11)
                replace_text(p, "<DeputyPosition>", group_commission.commission_member.position, 11)
                p = document.paragraphs[34]
                replace_text(p, "<DeputyFIO>", deputyFIO, 11)
            case "Член комиссии":
                p = document.paragraphs[24]
                memberFIO = f"{group_commission.commission_member.surname} {group_commission.commission_member.name[0]}. {group_commission.commission_member.patronymic[0]}."
                replace_text(p, "<MemberFIO>", memberFIO, 11)
                replace_text(p, "<MemberPosition>", group_commission.commission_member.position, 11)
                p = document.paragraphs[36]
                replace_text(p, "<MemberFIO>", memberFIO, 11)
            case "Секретарь":
                p = document.paragraphs[25]
                secretaryFIO = f"{group_commission.commission_member.surname} {group_commission.commission_member.name[0]}. {group_commission.commission_member.patronymic[0]}."
                replace_text(p, "<SecretaryFIO>", secretaryFIO, 11)
                p = document.paragraphs[38]
                replace_text(p, "<SecretaryFIO>", secretaryFIO, 11)
    
    p = document.paragraphs[26]
    match group.course.course_type:
        case "Профессиональная переподготовка":
            replace_text(p, "<CourseType>", "программе профессиональной подготовки", 11)
        case "Курсы повышения квалификации КПК":
            replace_text(p, "<CourseType>", "программе повышения квалификации КПК", 11)            
        case "Общеобразовательные программы для детей и взрослых":
            replace_text(p, "<CourseType>", "общеобразовательной программе для детей и взрослых", 11)            
        case "Профессиональное обучение":
            replace_text(p, "<CourseType>", "программе профессионального обучения", 11)
    replace_text(p, "<CourseName>", group.course.course_name, 11)
    
    p = document.paragraphs[27]
    begin_date_str = group.begin_date.strftime('%d.%m.%Y г.')
    finish_date_str = group.finish_date.strftime('%d.%m.%Y г.')
    print(p.text)
    replace_text(p, "<BeginDate>", f"{begin_date_str}", 11)
    replace_text(p, "<FinishDate>", f"{finish_date_str}", 11)
    
    table = document.tables[0]
    students = []
    if student_groups:
        students = [f"{student_group.student.surname} {student_group.student.name} {student_group.student.patronymic}"
                    for student_group in student_groups]
    
        students.sort(reverse=True)

        for row in table.rows[3:]:
            while students:
                row.cells[1].text = students.pop()
                p = row.cells[1].paragraphs[0]
                for run in p.runs:
                    run.font.name = "Cambria"
                    run.font.size = Pt(10)
                
                
    for row in table.rows:
        for cell in row.cells:
            if "<CourseName>" in cell.text:
                cell.text = group.course.course_name
                p = cell.paragraphs[0]
                set_font(p, 11)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    safe_group_name = sanitize_filename(group.name) 
    
    response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="protocol_{safe_group_name}.docx"'
    
    return response


def get_exam_sheet_docx(group_id, doc_number, date, order):
    group = Group.objects.get(id=group_id)
    expelled_students = StudentExpulsion.objects.filter(group=group).exclude(reason="Завершение курса").values_list("student_id", flat=True)
    student_groups = StudentGroup.objects.filter(group=group).exclude(student_id__in=expelled_students).select_related("student").order_by("student__surname")
    document = docx.Document('dpo/docs_patterns/exam_sheet.docx')

    date_obj = datetime.strptime(date, '%Y-%m-%d')
    formatted_date = date_obj.strftime('«%d» %B %Y г.')
    
    p = document.paragraphs[0]
    order_date = order.date.strftime('%d.%m.%Y')
    replace_text(p, "<OrderDate>", f"{order_date}", 11)
    replace_text(p, "<OrderNumber>", f"{order.number}", 11)
    
    p = document.paragraphs[5]
    replace_text(p, "<Number>", f"{doc_number}", 12)
    for run in p.runs:
        run.bold = True
    
    p = document.paragraphs[8]
    replace_text(p, "<CourseType>", f"{group.course.course_type}", 12)
    
    p = document.paragraphs[10]
    replace_text(p, "<CourseName>", f"{group.course.course_name}", 12)
    
    p = document.paragraphs[11]
    replace_text(p, "<Date>", f"{formatted_date}", 12)
    
    p = document.paragraphs[13]
    replace_text(p, "<StudentCount>", f"{len(student_groups)}", 11)

    table = document.tables[0]
    students = []
    if student_groups:
        students = [f"{student_group.student.surname} {student_group.student.name} {student_group.student.patronymic}"
                    for student_group in student_groups]
    
        students.sort(reverse=True)

        i = 1
        while students:
            new_row = table.add_row()
            new_row.cells[0].text = f"  {i}."
            p = new_row.cells[0].paragraphs[0]
            set_font(p, 12)
            
            new_row.cells[1].text = f"  {students.pop()}"
            p = new_row.cells[1].paragraphs[0]
            set_font(p, 10)
            
        new_row = table.add_row()
        new_row.cells[0].merge(new_row.cells[1])
        new_row.cells[0].text = "Средний балл итогового экзамена  "
        p = new_row.cells[0].paragraphs[0]
        set_font(p, 10)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(cell)
        row.height = None
        
    for group_commission in GroupCommission.objects.filter(group=group):
        match group_commission.role:
            case "Председатель комиссии":
                p = document.paragraphs[16]
                chairmanFIO = f"{group_commission.commission_member.surname} {group_commission.commission_member.name[0]}. {group_commission.commission_member.patronymic[0]}."
                replace_text(p, "<ChairmanFIO>", chairmanFIO, 11)
            case "Заместитель председателя комиссии":
                p = document.paragraphs[18]
                deputyFIO = f"{group_commission.commission_member.surname} {group_commission.commission_member.name[0]}. {group_commission.commission_member.patronymic[0]}."
                replace_text(p, "<DeputyFIO>", deputyFIO, 11)
            case "Член комиссии":
                p = document.paragraphs[20]
                memberFIO = f"{group_commission.commission_member.surname} {group_commission.commission_member.name[0]}. {group_commission.commission_member.patronymic[0]}."
                replace_text(p, "<MemberFIO>", memberFIO, 11)
            case "Секретарь":
                p = document.paragraphs[22]
                secretaryFIO = f"{group_commission.commission_member.surname} {group_commission.commission_member.name[0]}. {group_commission.commission_member.patronymic[0]}."
                replace_text(p, "<SecretaryFIO>", secretaryFIO, 11)
      
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    safe_group_name = sanitize_filename(group.name) 
    
    response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="exam_sheet_{safe_group_name}.docx"'
    
    return response


def get_expel_group_order_docx(group_id, doc_number, date, protocol):
    group = Group.objects.get(id=group_id)
    expelled_students = StudentExpulsion.objects.filter(reason="Завершение курса", group=group)
    student_groups = StudentGroup.objects.filter(group=group)
    students_ed_kind = []
    
    for student_group in student_groups:
        for expelled_student in expelled_students:
            if expelled_student.student.id == student_group.student.id:
                students_ed_kind.append((student_group.student, student_group.ed_kind))
    
    document = docx.Document('dpo/docs_patterns/expel_group.docx')

    table1 = document.tables[0]
    p = table1.cell(1, 0).paragraphs[0]
    table1.cell(1, 0).text = doc_number
    table1.cell(1, 1).text = date
    for cell in table1.rows[1].cells:
        for run in cell.paragraphs[0].runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = document.paragraphs[9]
    print(p.text, "p[9]")
    protocol_date = protocol.date.strftime('%d.%m.%Y г.')
    replace_text(p, "<ProtocolDate>", f"{protocol_date}", 11)
    replace_text(p, "<ProtocolNumber>", f"{protocol.number}", 11)
    
    p = document.paragraphs[12]
    print(p.text, "p[12]")
    match group.course.course_type:
        case "Профессиональная переподготовка":
            replace_text(p, "<CourseType>", "программе профессиональной подготовки", 11)
        case "Курсы повышения квалификации КПК":
            replace_text(p, "<CourseType>", "программе повышения квалификации КПК", 11)            
        case "Общеобразовательные программы для детей и взрослых":
            replace_text(p, "<CourseType>", "общеобразовательной программе для детей и взрослых", 11)            
        case "Профессиональное обучение":
            replace_text(p, "<CourseType>", "программе профессионального обучения", 11)
    replace_text(p, "<CourseName>", group.course.course_name, 11)

    p = document.paragraphs[14]
    print(p.text, "p[14]")
    replace_text(p, "<GroupName>", group.name, 12)
    for run in p.runs:
        run.bold = True
        
    p = document.paragraphs[15]
    print(p.text, "p[15]")
    replace_text(p, "<Price>", f"{group.course.price_b} руб., {group.course.price_vb}", 11)

    table2 = document.tables[1]
    
    i = 1
    for student_ed_kind in students_ed_kind:
        row = table2.add_row()
        row.cells[0].text = f"{i}."
        row.cells[1].text = f"{student_ed_kind[0].surname} {student_ed_kind[0].name} {student_ed_kind[0].patronymic}"
        if student_ed_kind[1] == "Внебюджет":
            row.cells[2].text = "ВБС"
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
        i += 1
        
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    safe_group_name = sanitize_filename(group.name) 
    
    response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="expel_group_order_{safe_group_name}.docx"'
    
    return response